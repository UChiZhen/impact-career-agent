"""DOCX resume rendering adapted from the legacy auto_resume project."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from career_agent.core import CandidateProfile


FONT_NAME = "Times New Roman"
PAGE_WIDTH = 11900
PAGE_HEIGHT = 16840
MARGIN_TOP = 522
MARGIN_RIGHT = 720
MARGIN_BOTTOM = 816
MARGIN_LEFT = 720
TAB_POSITION = 10460


def generate_resume_docx(
    tailored_data: dict[str, Any],
    candidate: CandidateProfile,
    output_path: Path,
) -> Path:
    """Generate a tailored resume DOCX from structured resume data."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Twips
    except ImportError as exc:
        raise RuntimeError(
            'DOCX rendering requires python-docx. Install "impact-career-agent[documents]".'
        ) from exc

    master_resume = candidate.master_resume or {}
    personal = {
        "name": candidate.name,
        "email": "",
        "phone": "",
        "location": candidate.location,
        **as_dict(master_resume.get("personal")),
    }
    education = normalize_education(master_resume.get("education") or candidate.education)

    def set_font(run, size, bold: bool = False):
        run.font.name = FONT_NAME
        run.font.size = size
        run.font.bold = bold
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), FONT_NAME)

    def set_spacing(paragraph, before=None, after=None, line=None, line_rule=None):
        ppr = paragraph._p.get_or_add_pPr()
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        if before is not None:
            spacing.set(qn("w:before"), str(before))
        if after is not None:
            spacing.set(qn("w:after"), str(after))
        if line is not None:
            spacing.set(qn("w:line"), str(line))
        if line_rule is not None:
            spacing.set(qn("w:lineRule"), line_rule)

    def add_right_tab(paragraph):
        ppr = paragraph._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(TAB_POSITION))
        tabs.append(tab)
        ppr.append(tabs)

    def add_bottom_border(paragraph):
        ppr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        pbdr.append(bottom)
        ppr.append(pbdr)

    def add_header(title: str):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(title)
        set_font(run, Pt(11), bold=True)
        add_bottom_border(paragraph)
        set_spacing(paragraph, after=72)

    def add_name():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(personal.get("name") or candidate.name)
        set_font(run, Pt(14), bold=True)

    def add_contact():
        values = [
            personal.get("email", ""),
            personal.get("phone", ""),
            personal.get("location", ""),
        ]
        contact_text = " | ".join(value for value in values if value)
        if not contact_text:
            return
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(contact_text)
        set_font(run, Pt(10))
        set_spacing(paragraph, after=240)

    def add_company_line(name: str, dates: str):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_right_tab(paragraph)
        set_spacing(paragraph, before=0, line=240, line_rule="auto")
        run = paragraph.add_run(name)
        set_font(run, Pt(10), bold=True)
        date_run = paragraph.add_run(f"\t{dates}")
        set_font(date_run, Pt(10), bold=True)

    def add_role_line(role: str, location: str):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_right_tab(paragraph)
        set_spacing(paragraph, before=0, line=240, line_rule="auto")
        run = paragraph.add_run(role)
        set_font(run, Pt(10))
        location_run = paragraph.add_run(f"\t{location}")
        set_font(location_run, Pt(10))

    def add_bullet(text: str, is_last: bool = False, section_gap: int = 120):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Twips(220)
        paragraph.paragraph_format.first_line_indent = Twips(-120)
        set_spacing(paragraph, before=0, line=240, line_rule="auto")
        if is_last:
            set_spacing(paragraph, after=section_gap)
        run = paragraph.add_run(text)
        set_font(run, Pt(10))

    def add_entry(entry: dict[str, Any], *, is_last_block: bool, section_gap: int):
        name = str(entry.get("company") or entry.get("org") or entry.get("name") or "")
        add_company_line(name, str(entry.get("dates", "")))
        add_role_line(str(entry.get("role", "")), str(entry.get("location", "")))
        bullets = [str(item) for item in entry.get("bullets", []) if str(item).strip()]
        for index, bullet in enumerate(bullets):
            add_bullet(
                bullet,
                is_last=is_last_block and index == len(bullets) - 1,
                section_gap=section_gap,
            )

    density = estimate_content_density(tailored_data)
    section_gap = 240 if density == "loose" else 120

    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH)
    section.page_height = Twips(PAGE_HEIGHT)
    section.top_margin = Twips(MARGIN_TOP)
    section.right_margin = Twips(MARGIN_RIGHT)
    section.bottom_margin = Twips(MARGIN_BOTTOM)
    section.left_margin = Twips(MARGIN_LEFT)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    add_name()
    add_contact()

    summary_text = str(tailored_data.get("summary_text", "")).strip()
    if summary_text:
        add_header("SUMMARY")
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(summary_text)
        set_font(run, Pt(10))
        set_spacing(paragraph, after=section_gap)

    if education:
        add_header("EDUCATION")
        for index, item in enumerate(education):
            add_company_line(item["school"], item["dates"])
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_right_tab(paragraph)
            set_spacing(paragraph, before=0, line=240, line_rule="auto")
            if index == len(education) - 1:
                set_spacing(paragraph, after=section_gap)
            degree_run = paragraph.add_run(item["degree"])
            set_font(degree_run, Pt(10))
            location_run = paragraph.add_run(f"\t{item['location']}")
            set_font(location_run, Pt(10))

    work_entries = [
        as_dict(entry) for entry in tailored_data.get("work_experience", []) if as_dict(entry)
    ]
    if work_entries:
        add_header(str(tailored_data.get("work_experience_header", "WORK EXPERIENCE")))
        for index, entry in enumerate(work_entries):
            add_entry(entry, is_last_block=index == len(work_entries) - 1, section_gap=section_gap)

    combined_entries = [
        as_dict(entry) for entry in tailored_data.get("combined_section", []) if as_dict(entry)
    ]
    if combined_entries:
        add_header(str(tailored_data.get("combined_section_header", "SELECTED PROJECTS")))
        for index, entry in enumerate(combined_entries):
            add_entry(entry, is_last_block=index == len(combined_entries) - 1, section_gap=section_gap)

    skills = normalize_skills(tailored_data.get("skills") or master_resume.get("skills"))
    if skills:
        add_header("SKILLS")
        for index, skill in enumerate(skills):
            paragraph = doc.add_paragraph()
            set_spacing(paragraph, before=0, after=0 if index == len(skills) - 1 else 120)
            label_run = paragraph.add_run(f"{skill['label']}: ")
            set_font(label_run, Pt(10), bold=True)
            value_run = paragraph.add_run(skill["value"])
            set_font(value_run, Pt(10))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def estimate_content_density(tailored_data: dict[str, Any]) -> str:
    """Estimate whether the resume needs tight spacing to stay compact."""
    chars_per_line = 90
    lines = 8
    lines += max(1, len(str(tailored_data.get("summary_text", ""))) // chars_per_line + 1)
    for key in ("work_experience", "combined_section"):
        entries = tailored_data.get(key, []) or []
        lines += len(entries) * 2
        for entry in entries:
            for bullet in as_dict(entry).get("bullets", []) or []:
                lines += max(1, len(str(bullet)) // chars_per_line + 1)
    lines += len(tailored_data.get("skills", []) or [])
    return "loose" if lines < 39 else "tight"


def normalize_education(value: Any) -> list[dict[str, str]]:
    """Normalize master-resume education rows."""
    rows = []
    for item in value if isinstance(value, list) else [value]:
        if isinstance(item, str):
            rows.append({"school": item, "degree": "", "location": "", "dates": ""})
        elif isinstance(item, dict):
            rows.append(
                {
                    "school": str(item.get("school") or item.get("name") or ""),
                    "degree": str(item.get("degree") or item.get("details") or ""),
                    "location": str(item.get("location", "")),
                    "dates": str(item.get("dates", "")),
                }
            )
    return [row for row in rows if row["school"] or row["degree"]]


def normalize_skills(value: Any) -> list[dict[str, str]]:
    """Normalize skills to label/value rows."""
    if isinstance(value, dict):
        rows = []
        for key, item in value.items():
            if isinstance(item, list):
                skill_value = ", ".join(str(skill) for skill in item)
            else:
                skill_value = str(item)
            if skill_value.strip():
                rows.append({"label": str(key).replace("_", " ").title(), "value": skill_value})
        return rows

    rows = []
    for item in value if isinstance(value, list) else [value]:
        if isinstance(item, dict):
            label = item.get("label") or item.get("name") or item.get("category") or "Skills"
            skill_value = item.get("value") or item.get("content") or item.get("skills") or ""
            if isinstance(skill_value, list):
                skill_value = ", ".join(str(skill) for skill in skill_value)
            if str(skill_value).strip():
                rows.append({"label": str(label), "value": str(skill_value)})
        elif item:
            rows.append({"label": "Skills", "value": str(item)})
    return rows


def as_dict(value: Any) -> dict[str, Any]:
    """Return a dict or an empty dict."""
    return value if isinstance(value, dict) else {}
