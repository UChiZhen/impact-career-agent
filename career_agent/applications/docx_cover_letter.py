"""DOCX cover-letter rendering adapted from the legacy auto_resume project."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from career_agent.core import CandidateProfile


FONT_NAME = "Times New Roman"
PAGE_WIDTH = 11900
PAGE_HEIGHT = 16840
MARGIN_TOP = 816
MARGIN_RIGHT = 1225
MARGIN_BOTTOM = 816
MARGIN_LEFT = 1225
FIRST_LINE_INDENT = 357
LINE_SPACING = 276


def generate_cover_letter_docx(
    cover_letter_data: dict[str, Any],
    candidate: CandidateProfile,
    output_path: Path,
) -> Path:
    """Generate a tailored cover-letter DOCX."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Twips
    except ImportError as exc:
        raise RuntimeError(
            'DOCX rendering requires python-docx. Install "impact-career-agent[documents]".'
        ) from exc

    master_resume = candidate.master_resume or {}
    personal = {
        "name": candidate.name,
        "email": "",
        "phone": "",
        "location": candidate.location,
        **(master_resume.get("personal") if isinstance(master_resume.get("personal"), dict) else {}),
    }
    paragraphs = [
        str(paragraph)
        for paragraph in cover_letter_data.get("paragraphs", [])
        if str(paragraph).strip()
    ]
    body_size = choose_body_font_size(paragraphs)

    def set_font(run, size, bold: bool = False):
        run.font.name = FONT_NAME
        run.font.size = size
        run.font.bold = bold
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), FONT_NAME)

    def set_spacing(paragraph, after=72, line=LINE_SPACING):
        ppr = paragraph._p.get_or_add_pPr()
        spacing = ppr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            ppr.append(spacing)
        spacing.set(qn("w:after"), str(after))
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")

    def set_first_line_indent(paragraph):
        ppr = paragraph._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), str(FIRST_LINE_INDENT))
        ppr.append(ind)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(PAGE_WIDTH)
    section.page_height = Twips(PAGE_HEIGHT)
    section.top_margin = Twips(MARGIN_TOP)
    section.right_margin = Twips(MARGIN_RIGHT)
    section.bottom_margin = Twips(MARGIN_BOTTOM)
    section.left_margin = Twips(MARGIN_LEFT)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = body_size
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(name.add_run(personal.get("name") or candidate.name), Pt(14), bold=True)

    contact_values = [personal.get("email", ""), personal.get("phone", ""), personal.get("location", "")]
    contact_text = " | ".join(value for value in contact_values if value)
    if contact_text:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(contact.add_run(contact_text), Pt(10))

    today = date.today()
    dated = doc.add_paragraph()
    dated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(dated, after=72)
    set_font(dated.add_run(format_cover_letter_date(today)), Pt(10))

    greeting = doc.add_paragraph()
    set_spacing(greeting)
    set_font(greeting.add_run(cover_letter_data.get("greeting", "Dear Hiring Manager,")), body_size)

    for paragraph_text in paragraphs:
        paragraph = doc.add_paragraph()
        set_spacing(paragraph)
        set_first_line_indent(paragraph)
        set_font(paragraph.add_run(paragraph_text), body_size)

    closing = doc.add_paragraph()
    set_spacing(closing)
    set_font(closing.add_run(cover_letter_data.get("closing", "Best regards,")), body_size)

    signature = doc.add_paragraph()
    set_spacing(signature, after=0)
    set_font(signature.add_run(cover_letter_data.get("signature") or candidate.name), body_size)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def choose_body_font_size(paragraphs: list[str]):
    """Choose cover-letter body font size based on content length."""
    from docx.shared import Pt

    total_words = sum(len(paragraph.split()) for paragraph in paragraphs)
    if total_words < 300:
        return Pt(14)
    if total_words < 420:
        return Pt(13)
    return Pt(12)


def format_cover_letter_date(value: date) -> str:
    """Format date for cover letters without superscript complexity."""
    return f"{value.strftime('%B')} {value.day}, {value.year}"
